import os
from datetime import datetime
from typing import Dict, List, Tuple
try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore[assignment,misc]

from .preset import Preset
from .layer1 import analyze_layer1_text
from .findings import Finding


def scrub_docx(
    input_path: str,
    output_path: str,
    preset: Preset,
    language: str,
    file_id: str,
) -> Tuple[List[Finding], Dict[str, int]]:

    # Select analysis function based on preset layer
    if preset.layer == 2:
        from .layer2_candle import analyze_layer2_text as _analyze
    elif preset.layer == 3:
        from .layer3_presidio import analyze_layer3_text as _analyze
    else:
        _analyze = analyze_layer1_text

    doc = Document(input_path)
    all_findings: List[Finding] = []
    summary: Dict[str, int] = {}
    filename = os.path.basename(input_path)

    def process(text: str, location: str) -> str:
        redacted, findings, local = _analyze(text, preset, language)
        for f in findings:
            f.file_id = file_id
            f.original_filename = filename
            f.page_or_location = location
        all_findings.extend(findings)
        for k, v in local.items():
            summary[k] = summary.get(k, 0) + v
        return redacted

    # ── Body paragraphs ───────────────────────────────────────────────────────
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            for run in p.runs:
                if run.text:
                    run.text = process(run.text, f"paragraph {i + 1}")

    # ── Tables ────────────────────────────────────────────────────────────────
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                if cell.text.strip():
                    for p in cell.paragraphs:
                        for run in p.runs:
                            if run.text:
                                run.text = process(
                                    run.text,
                                    f"table {ti + 1}, row {ri + 1}, cell {ci + 1}",
                                )

    # ── Headers and footers ───────────────────────────────────────────────────
    for section in doc.sections:
        for hf_name, hf_obj in [
            ("header", section.header),
            ("footer", section.footer),
            ("first_page_header", section.first_page_header),
            ("first_page_footer", section.first_page_footer),
            ("even_page_header", section.even_page_header),
            ("even_page_footer", section.even_page_footer),
        ]:
            if hf_obj is None:
                continue
            try:
                for p in hf_obj.paragraphs:
                    for run in p.runs:
                        if run.text:
                            run.text = process(run.text, hf_name)
            except Exception:
                pass  # Some header/footer objects are unlinked; skip safely

    # ── Document core properties (metadata) ───────────────────────────────────
    # Clear fields that may contain author/reviewer PII.
    cp = doc.core_properties
    cp.author = "Anonymized"
    cp.last_modified_by = "Anonymized"
    cp.title = ""
    cp.subject = ""
    cp.description = ""
    cp.keywords = ""
    cp.category = ""
    cp.comments = ""
    cp.content_status = ""
    # Revision counter reset and timestamps cleared
    try:
        cp.revision = 1
    except Exception:
        pass
    try:
        now = datetime(1970, 1, 1)
        cp.created = now
        cp.modified = now
        cp.last_printed = now
    except Exception:
        pass

    # ── Remove tracked changes (revisions) ───────────────────────────────────
    # Tracked changes embed original text, which would leak the PII we just removed.
    body = doc.element.body
    for tag in ("w:ins", "w:del", "w:rPrChange", "w:pPrChange", "w:sectPrChange"):
        ns, local = tag.split(":")
        for elem in body.findall(f".//{{{doc.element.nsmap.get(ns, 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')}}}{local}"):
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)

    # ── Remove comments ───────────────────────────────────────────────────────
    try:
        comments_part = doc.part.package.part_related_by(
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
        )
        comments_xml = comments_part._element
        for comment in comments_xml.findall(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comment"
        ):
            comments_xml.remove(comment)
    except Exception:
        pass  # No comments part present

    doc.save(output_path)
    return all_findings, summary
